import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def save_dict_to_file(dictionary, file_path, file_name, file_type, keys=None):
    """
    将字典保存为本地文件，支持 Word、PDF 和 Markdown 格式。
    
    参数:
        dictionary (dict): 要保存的字典
        file_path (str): 文件保存路径
        file_name (str): 文件名
        file_type (str): 文件类型，支持 'word'、'pdf' 或 'markdown'
        keys (list, optional): 要保存的键列表。如果为 None，则保存整个字典。
    """
    # 如果指定了 keys，则过滤字典
    if keys:
        dictionary = {key: dictionary[key] for key in keys if key in dictionary}
    
    # 确保文件路径存在
    os.makedirs(file_path, exist_ok=True)
    
    # 根据文件类型选择保存方式
    if file_type.lower() == 'word':
        # 保存为 Word 文件
        doc = Document()
        for key, value in dictionary.items():
            doc.add_heading(key, level=1)
            doc.add_paragraph(str(value))
        doc.save(os.path.join(file_path, f"{file_name}.docx"))
        print(f"字典已保存为 Word 文件: {file_name}.docx")
    
    elif file_type.lower() == 'pdf':
        # 保存为 PDF 文件
        pdf_path = os.path.join(file_path, f"{file_name}.pdf")
        c = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter
        y = height - 100
        
        for key, value in dictionary.items():
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y, key)
            y -= 20
            c.setFont("Helvetica", 12)
            c.drawString(70, y, str(value))
            y -= 30
            if y < 50:
                c.showPage()
                y = height - 100
        
        c.save()
        print(f"字典已保存为 PDF 文件: {file_name}.pdf")
    
    elif file_type.lower() == 'markdown':
        # 保存为 Markdown 文件
        md_path = os.path.join(file_path, f"{file_name}.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            for key, value in dictionary.items():
                f.write(f"### {key}\n")
                f.write(f"{value}\n\n")
        print(f"字典已保存为 Markdown 文件: {file_name}.md")
    
    else:
        raise ValueError("不支持的文件类型，请选择 'word'、'pdf' 或 'markdown'")

# 示例用法
if __name__ == "__main__":
    my_dict = {
        "姓名": "张三",
        "年龄": 30,
        "职业": "程序员",
        "兴趣爱好": ["编程", "阅读", "旅行"]
    }
    
    # 保存指定键为 Word 文件
    save_dict_to_file(my_dict, "Results/output", "example_selected", "word", keys=["姓名", "职业"])
    
    # 保存指定键为 PDF 文件
    save_dict_to_file(my_dict, "Results/output", "example_selected", "pdf", keys=["年龄", "兴趣爱好"])
    
    # 保存指定键为 Markdown 文件
    save_dict_to_file(my_dict, "Results/output", "example_selected", "markdown", keys=["姓名", "兴趣爱好"])